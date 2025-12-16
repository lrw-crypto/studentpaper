import os
import re
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from io import BytesIO

app = Flask(__name__)

# --- 核心排版功能區 ---

def add_page_number(run):
    """
    在 python-docx 中插入頁碼需要操作底層 XML
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_chinese_font(run, size_pt=12):
    """
    強制設定中英文雙字型：
    中文：新細明體 (PMingLiU)
    英文：Times New Roman
    """
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')
    run.font.size = Pt(size_pt)
    run.font.color.rgb = None # 確保是黑色(預設)

def format_thesis(doc_stream, paper_title):
    doc = Document(doc_stream)
    
    # 1. 版面設定 (Page Layout) [cite: 323, 333]
    # 規定：A4, 上下左右邊界皆為 2公分
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 2. 設定預設樣式 (Normal Style) [cite: 332]
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.0 # 單行間距
    set_chinese_font(style, 12) # 預設全域字體設定

    # 3. 頁首與頁尾 (Header & Footer) [cite: 337]
    # 頁首：小論文篇名 (置中, 10pt)
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        header_para = header.paragraphs[0]
        header_para.clear()
    else:
        header_para = header.add_paragraph()
    
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = header_para.add_run(paper_title)
    set_chinese_font(h_run, 10)

    # 頁尾：頁碼 (置中, 10pt)
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        footer_para = footer.paragraphs[0]
        footer_para.clear()
    else:
        footer_para = footer.add_paragraph()
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_para.add_run()
    add_page_number(f_run)
    set_chinese_font(f_run, 10)

    # 4. 內容遍歷與智慧格式化
    is_reference_section = False
    
    # 六大標準標題強制對應表 (含英文) [cite: 339, 404-406]
    standard_headings = {
        # 中文標準
        "前言": "壹、前言",
        "文獻探討": "貳、文獻探討",
        "研究方法": "參、研究方法",
        "研究分析與結果": "肆、研究分析與結果",
        "研究結論與建議": "伍、研究結論與建議",
        "參考文獻": "陸、參考文獻",
        "引註資料": "陸、參考文獻",
        # 英文標準
        "Introduction": "I. Introduction",
        "LiteratureReview": "II. Literature Review",
        "ResearchMethods": "III. Research Methods",
        "AnalysisandResults": "IV. Analysis and Results",
        "ConclusionandSuggestions": "V. Conclusion and Suggestions",
        "References": "VI. References"
    }

    # Regex 定義
    # 中文層級
    regex_h1_zh = re.compile(r'^[壹貳參肆伍陸]、')
    regex_h2_zh = re.compile(r'^[一二三四五六七八九十]+、')
    regex_h3_zh = re.compile(r'^\（[一二三四五六七八九十]+\）')
    regex_h4_zh = re.compile(r'^[０-９]+、')

    # 英文層級 Regex (支援英文小論文標準) [cite: 350-353]
    # Level 1: I. II. III.
    regex_h1_en = re.compile(r'^(I|II|III|IV|V|VI)\.\s*') 
    # Level 2: (I) (II)
    regex_h2_en = re.compile(r'^\((I|II|III|IV|V|VI)\)\s*')
    # Level 3: A. B.
    regex_h3_en = re.compile(r'^[A-Z]\.\s*')

    # 圖表計數器
    fig_count = 0
    table_count = 0
    # 支援英文圖表標題 (Figure/Table)
    regex_fig_caption = re.compile(r'^(圖|Figure)\s*([0-9]+|[一二三四五六七八九十]+)(.*)', re.IGNORECASE)
    regex_table_caption = re.compile(r'^(表|Table)\s*([0-9]+|[一二三四五六七八九十]+)(.*)', re.IGNORECASE)

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 跳過空白段落，但保留格式
        if not text:
            continue

        # 清理文字以便比對
        clean_text = text.replace(" ", "").replace("　", "")
        
        # --- 自動修正與強制六大標題 ---
        for key, correct_title in standard_headings.items():
            # 檢查是否包含關鍵字，且長度短（避免誤判內文）
            if key.lower() in clean_text.lower() and len(clean_text) < 30:
                # 排除已經正確的情況，若不正確則強制修正
                if clean_text != correct_title.replace(" ", ""):
                    para.clear()
                    run = para.add_run(correct_title)
                    set_chinese_font(run, 12)
                    run.bold = False # 規定標題不粗體 [cite: 328]
                    text = correct_title # 更新變數供後續邏輯使用
                    
                # 確保前言/Introduction 不強制換頁 (避免封面頁效果)
                if "前言" in correct_title or "Introduction" in correct_title:
                    para.paragraph_format.page_break_before = False
                    
                # 確保大標靠左不縮排 [cite: 334]
                para.paragraph_format.first_line_indent = Pt(0)
                break

        # --- 圖表編號自動修正 ---
        match_fig = regex_fig_caption.match(text)
        if match_fig:
            fig_count += 1
            # 判斷是中文還是英文前綴
            prefix = "Figure" if "Figure" in match_fig.group(1) or "figure" in match_fig.group(1) else "圖"
            new_caption = f"{prefix} {fig_count} {match_fig.group(3).strip()}"
            if text != new_caption:
                para.clear()
                run = para.add_run(new_caption)
                set_chinese_font(run, 12)
                run.bold = False
            # 規定：圖表標題均置於圖表上方置左 [cite: 358]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            continue

        match_table = regex_table_caption.match(text)
        if match_table:
            table_count += 1
            prefix = "Table" if "Table" in match_table.group(1) or "table" in match_table.group(1) else "表"
            new_caption = f"{prefix} {table_count} {match_table.group(3).strip()}"
            if text != new_caption:
                para.clear()
                run = para.add_run(new_caption)
                set_chinese_font(run, 12)
                run.bold = False
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            continue

        # --- 參考文獻區域判斷 ---
        # 同時偵測中文「陸、參考文獻」與英文「VI. References」
        is_ref_keyword = ("參考文獻" in text and text.startswith("陸")) or \
                         ("References" in text and (text.startswith("VI") or text.startswith("6")))
        
        if is_ref_keyword:
            is_reference_section = True
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.left_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                set_chinese_font(run, 12)
                run.bold = False
            continue

        # --- 處理「參考文獻」內容的 APA 格式 (懸掛縮排) ---
        if is_reference_section:
            para.paragraph_format.left_indent = Pt(24)
            para.paragraph_format.first_line_indent = Pt(-24)
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                is_italic = run.italic
                set_chinese_font(run, 12)
                run.italic = is_italic
                run.bold = False
            continue

        # --- 一般內文與標題層級處理 ---
        # 預設：首行縮排 2 字元 (約 24pt) [cite: 335]
        para.paragraph_format.first_line_indent = Pt(24) 
        para.paragraph_format.left_indent = Pt(0)
        
        is_heading = False
        
        # 中文標題判斷
        if regex_h1_zh.match(text) or regex_h2_zh.match(text): # 壹、 或 一、
            is_heading = True
            para.paragraph_format.first_line_indent = Pt(0)
        elif regex_h3_zh.match(text) or regex_h4_zh.match(text): # (一) 或 1、
            is_heading = True
            para.paragraph_format.first_line_indent = Pt(24) # 依規定，此層級通常有縮排

        # 英文標題判斷 - 解決內縮問題 [cite: 350-352]
        elif regex_h1_en.match(text): # I. II. -> 大標題不縮排
            is_heading = True
            para.paragraph_format.first_line_indent = Pt(0)
        elif regex_h2_en.match(text): # (I) (II) -> 視為次級標題，設為不縮排以區隔內文
            is_heading = True
            para.paragraph_format.first_line_indent = Pt(0)
        elif regex_h3_en.match(text): # A. B. -> 視為內文列表，維持縮排
            is_heading = True
            para.paragraph_format.first_line_indent = Pt(24)

        # 統一字體設定
        for run in para.runs:
            set_chinese_font(run, 12)
            run.bold = False 
            run.italic = False
            run.underline = False
            
    # 存入緩衝區
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

# --- 路由區 ---

@app.route('/', methods=['GET'])
def index():
    # 改為使用模板，請確保 templates/index.html 存在
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return '沒有上傳檔案', 400
    
    file = request.files['file']
    title = request.form.get('title', '小論文')
    
    if file.filename == '':
        return '未選擇檔案', 400
        
    if file and file.filename.endswith('.docx'):
        # 執行排版
        formatted_file = format_thesis(file, title)
        
        return send_file(
            formatted_file,
            as_attachment=True,
            download_name=f'已排版_{file.filename}',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        return '請上傳 .docx 格式的 Word 檔案', 400

if __name__ == '__main__':
    # 自動建立 templates 資料夾 (如果不存在)，防止錯誤
    if not os.path.exists('templates'):
        os.makedirs('templates', exist_ok=True)
        print("警告：templates 資料夾已自動建立，請確保 index.html 檔案已放入其中。")
        
    app.run(debug=True, port=5000)
